"""Document generation service for resumes and cover letters."""
import os
from typing import Dict, Optional, List
from datetime import datetime
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
from reportlab.lib import colors
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from services.gemini_service import GeminiService


class DocumentGenerator:
    """Generate professional resumes and cover letters."""

    def __init__(self):
        """Initialize document generator."""
        self.gemini = GeminiService()

    def generate_resume(
        self,
        user_data: Dict,
        format: str = 'pdf',
        template: str = 'modern'
    ) -> str:
        """
        Generate a professional resume from user data.

        Args:
            user_data: User profile and resume data
            format: Output format ('pdf' or 'docx')
            template: Resume template ('modern', 'classic', 'minimal')

        Returns:
            Path to generated resume file
        """
        if template == 'modern':
            return self._generate_modern_resume(user_data, format)
        elif template == 'classic':
            return self._generate_classic_resume(user_data, format)
        else:
            return self._generate_minimal_resume(user_data, format)

    def generate_cover_letter(
        self,
        user_data: Dict,
        job_data: Dict,
        format: str = 'pdf',
        custom_content: Optional[str] = None
    ) -> str:
        """
        Generate a tailored cover letter for a job application.

        Args:
            user_data: User profile data
            job_data: Job posting information
            format: Output format ('pdf' or 'docx')
            custom_content: Optional custom content to include

        Returns:
            Path to generated cover letter file
        """
        # Use AI to generate personalized cover letter content
        cover_letter_content = self._generate_cover_letter_content(
            user_data, job_data, custom_content
        )

        if format == 'pdf':
            return self._create_cover_letter_pdf(user_data, job_data, cover_letter_content)
        else:
            return self._create_cover_letter_docx(user_data, job_data, cover_letter_content)

    def _generate_modern_resume(self, user_data: Dict, format: str) -> str:
        """Generate modern style resume."""
        if format == 'pdf':
            return self._create_modern_resume_pdf(user_data)
        else:
            return self._create_modern_resume_docx(user_data)

    def _create_modern_resume_pdf(self, user_data: Dict) -> str:
        """Create modern resume in PDF format."""
        import tempfile

        # Create temp file
        temp_dir = tempfile.gettempdir()
        filename = f"resume_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        filepath = os.path.join(temp_dir, filename)

        # Create PDF
        doc = SimpleDocTemplate(filepath, pagesize=letter)
        story = []
        styles = getSampleStyleSheet()

        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#2c3e50'),
            spaceAfter=6,
            alignment=1  # Center
        )

        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#3498db'),
            spaceAfter=12,
            spaceBefore=12,
            borderWidth=1,
            borderColor=colors.HexColor('#3498db'),
            borderPadding=5
        )

        # Personal Info
        personal = user_data.get('personalInfo', {})
        name = personal.get('fullName', 'Your Name')
        story.append(Paragraph(name, title_style))

        contact_info = []
        if personal.get('email'):
            contact_info.append(personal['email'])
        if personal.get('phone'):
            contact_info.append(personal['phone'])
        if personal.get('location'):
            contact_info.append(personal['location'])

        if contact_info:
            contact_text = ' | '.join(contact_info)
            story.append(Paragraph(contact_text, styles['Normal']))

        story.append(Spacer(1, 0.2*inch))

        # Summary
        summary = user_data.get('summary', '')
        if summary:
            story.append(Paragraph('PROFESSIONAL SUMMARY', heading_style))
            story.append(Paragraph(summary, styles['Normal']))
            story.append(Spacer(1, 0.2*inch))

        # Work Experience
        experiences = user_data.get('experiences', [])
        if experiences:
            story.append(Paragraph('WORK EXPERIENCE', heading_style))
            for exp in experiences:
                title = exp.get('title', 'Position')
                company = exp.get('company', 'Company')
                dates = f"{exp.get('startDate', '')} - {exp.get('endDate', 'Present')}"

                story.append(Paragraph(f"<b>{title}</b> at {company}", styles['Normal']))
                story.append(Paragraph(f"<i>{dates}</i>", styles['Normal']))

                if exp.get('description'):
                    story.append(Paragraph(exp['description'], styles['Normal']))

                if exp.get('achievements'):
                    for achievement in exp['achievements']:
                        story.append(Paragraph(f"• {achievement}", styles['Normal']))

                story.append(Spacer(1, 0.1*inch))

        # Education
        education = user_data.get('education', [])
        if education:
            story.append(Paragraph('EDUCATION', heading_style))
            for edu in education:
                degree = edu.get('degree', 'Degree')
                institution = edu.get('institution', 'Institution')
                dates = f"{edu.get('startDate', '')} - {edu.get('endDate', '')}"

                story.append(Paragraph(f"<b>{degree}</b>", styles['Normal']))
                story.append(Paragraph(f"{institution} | {dates}", styles['Normal']))

                if edu.get('gpa'):
                    story.append(Paragraph(f"GPA: {edu['gpa']}", styles['Normal']))

                story.append(Spacer(1, 0.1*inch))

        # Skills
        skills = user_data.get('skills', [])
        if skills:
            story.append(Paragraph('SKILLS', heading_style))
            skills_text = ', '.join(skills[:20])  # Top 20 skills
            story.append(Paragraph(skills_text, styles['Normal']))

        # Build PDF
        doc.build(story)
        return filepath

    def _create_modern_resume_docx(self, user_data: Dict) -> str:
        """Create modern resume in DOCX format."""
        import tempfile

        # Create temp file
        temp_dir = tempfile.gettempdir()
        filename = f"resume_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(temp_dir, filename)

        # Create document
        doc = Document()

        # Personal Info
        personal = user_data.get('personalInfo', {})
        name = personal.get('fullName', 'Your Name')

        heading = doc.add_heading(name, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Contact info
        contact_info = []
        if personal.get('email'):
            contact_info.append(personal['email'])
        if personal.get('phone'):
            contact_info.append(personal['phone'])
        if personal.get('location'):
            contact_info.append(personal['location'])

        if contact_info:
            contact_para = doc.add_paragraph(' | '.join(contact_info))
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Summary
        summary = user_data.get('summary', '')
        if summary:
            doc.add_heading('PROFESSIONAL SUMMARY', 1)
            doc.add_paragraph(summary)

        # Work Experience
        experiences = user_data.get('experiences', [])
        if experiences:
            doc.add_heading('WORK EXPERIENCE', 1)
            for exp in experiences:
                title = exp.get('title', 'Position')
                company = exp.get('company', 'Company')
                dates = f"{exp.get('startDate', '')} - {exp.get('endDate', 'Present')}"

                job_para = doc.add_paragraph()
                job_para.add_run(f"{title} at {company}").bold = True
                doc.add_paragraph(dates).italic = True

                if exp.get('description'):
                    doc.add_paragraph(exp['description'])

                if exp.get('achievements'):
                    for achievement in exp['achievements']:
                        doc.add_paragraph(achievement, style='List Bullet')

        # Education
        education = user_data.get('education', [])
        if education:
            doc.add_heading('EDUCATION', 1)
            for edu in education:
                degree = edu.get('degree', 'Degree')
                institution = edu.get('institution', 'Institution')
                dates = f"{edu.get('startDate', '')} - {edu.get('endDate', '')}"

                edu_para = doc.add_paragraph()
                edu_para.add_run(degree).bold = True
                doc.add_paragraph(f"{institution} | {dates}")

                if edu.get('gpa'):
                    doc.add_paragraph(f"GPA: {edu['gpa']}")

        # Skills
        skills = user_data.get('skills', [])
        if skills:
            doc.add_heading('SKILLS', 1)
            doc.add_paragraph(', '.join(skills[:20]))

        # Save document
        doc.save(filepath)
        return filepath

    def _generate_cover_letter_content(
        self,
        user_data: Dict,
        job_data: Dict,
        custom_content: Optional[str] = None
    ) -> str:
        """Generate personalized cover letter content using AI."""

        # Prepare data for AI
        personal = user_data.get('personalInfo', {})
        experiences = user_data.get('experiences', [])
        skills = user_data.get('skills', [])

        job_title = job_data.get('title', 'the position')
        company_name = job_data.get('company', 'your company')
        job_description = job_data.get('description', '')

        prompt = f"""Write a professional cover letter for the following job application:

Job Title: {job_title}
Company: {company_name}
Job Description: {job_description[:500]}

Applicant Information:
Name: {personal.get('fullName', 'the applicant')}
Skills: {', '.join(skills[:10])}
Recent Experience: {experiences[0].get('title', '') if experiences else 'N/A'} at {experiences[0].get('company', '') if experiences else 'N/A'}

Requirements:
1. Write a compelling opening paragraph that shows enthusiasm for the role
2. Highlight relevant skills and experiences that match the job requirements
3. Provide 2-3 specific examples of achievements that demonstrate value
4. Express genuine interest in the company and role
5. Include a strong closing paragraph with call to action
6. Keep it professional, confident, and concise (300-400 words)
7. Do NOT include address blocks or date (just the body content)

{f'Additional context to include: {custom_content}' if custom_content else ''}

Write ONLY the body paragraphs of the cover letter (no addresses, no date, no signature block):
"""

        try:
            response = self.gemini.generate_text(prompt)
            return response
        except Exception as e:
            # Fallback to template if AI fails
            return self._generate_template_cover_letter(user_data, job_data)

    def _generate_template_cover_letter(
        self,
        user_data: Dict,
        job_data: Dict
    ) -> str:
        """Generate a template-based cover letter as fallback."""
        personal = user_data.get('personalInfo', {})
        name = personal.get('fullName', 'the applicant')
        job_title = job_data.get('title', 'this position')
        company_name = job_data.get('company', 'your company')

        skills = user_data.get('skills', [])
        top_skills = ', '.join(skills[:5]) if skills else 'my diverse skill set'

        return f"""I am writing to express my strong interest in the {job_title} position at {company_name}. With my background in {top_skills}, I am excited about the opportunity to contribute to your team.

Throughout my career, I have developed a strong foundation in the skills required for this role. My experience has equipped me with the ability to tackle complex challenges and deliver results that drive business success. I am particularly drawn to {company_name} because of your commitment to excellence and innovation in the industry.

I am confident that my technical expertise and passion for continuous learning make me an ideal candidate for this position. I would welcome the opportunity to discuss how my skills and experiences align with your team's needs.

Thank you for considering my application. I look forward to the possibility of discussing this opportunity further."""

    def _create_cover_letter_pdf(
        self,
        user_data: Dict,
        job_data: Dict,
        content: str
    ) -> str:
        """Create cover letter in PDF format."""
        import tempfile

        temp_dir = tempfile.gettempdir()
        filename = f"cover_letter_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        filepath = os.path.join(temp_dir, filename)

        doc = SimpleDocTemplate(filepath, pagesize=letter)
        story = []
        styles = getSampleStyleSheet()

        # Header with contact info
        personal = user_data.get('personalInfo', {})
        name = personal.get('fullName', 'Your Name')

        header_style = ParagraphStyle(
            'Header',
            parent=styles['Normal'],
            fontSize=11,
            alignment=0
        )

        story.append(Paragraph(name, header_style))
        if personal.get('email'):
            story.append(Paragraph(personal['email'], header_style))
        if personal.get('phone'):
            story.append(Paragraph(personal['phone'], header_style))

        story.append(Spacer(1, 0.3*inch))

        # Date
        today = datetime.now().strftime('%B %d, %Y')
        story.append(Paragraph(today, styles['Normal']))
        story.append(Spacer(1, 0.2*inch))

        # Company info
        company_name = job_data.get('company', 'Hiring Manager')
        job_title = job_data.get('title', 'Position')

        story.append(Paragraph(f"Hiring Manager", styles['Normal']))
        story.append(Paragraph(company_name, styles['Normal']))
        story.append(Spacer(1, 0.2*inch))

        # Salutation
        story.append(Paragraph(f"Dear Hiring Manager,", styles['Normal']))
        story.append(Spacer(1, 0.1*inch))

        # Content paragraphs
        for paragraph in content.split('\n\n'):
            if paragraph.strip():
                story.append(Paragraph(paragraph.strip(), styles['Normal']))
                story.append(Spacer(1, 0.1*inch))

        # Closing
        story.append(Spacer(1, 0.2*inch))
        story.append(Paragraph("Sincerely,", styles['Normal']))
        story.append(Spacer(1, 0.3*inch))
        story.append(Paragraph(name, styles['Normal']))

        doc.build(story)
        return filepath

    def _create_cover_letter_docx(
        self,
        user_data: Dict,
        job_data: Dict,
        content: str
    ) -> str:
        """Create cover letter in DOCX format."""
        import tempfile

        temp_dir = tempfile.gettempdir()
        filename = f"cover_letter_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(temp_dir, filename)

        doc = Document()

        # Header with contact info
        personal = user_data.get('personalInfo', {})
        name = personal.get('fullName', 'Your Name')

        doc.add_paragraph(name)
        if personal.get('email'):
            doc.add_paragraph(personal['email'])
        if personal.get('phone'):
            doc.add_paragraph(personal['phone'])

        doc.add_paragraph()  # Blank line

        # Date
        today = datetime.now().strftime('%B %d, %Y')
        doc.add_paragraph(today)
        doc.add_paragraph()

        # Company info
        company_name = job_data.get('company', 'Hiring Manager')
        doc.add_paragraph('Hiring Manager')
        doc.add_paragraph(company_name)
        doc.add_paragraph()

        # Salutation
        doc.add_paragraph('Dear Hiring Manager,')
        doc.add_paragraph()

        # Content paragraphs
        for paragraph in content.split('\n\n'):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())

        # Closing
        doc.add_paragraph()
        doc.add_paragraph('Sincerely,')
        doc.add_paragraph()
        doc.add_paragraph(name)

        doc.save(filepath)
        return filepath

    def _generate_classic_resume(self, user_data: Dict, format: str) -> str:
        """Generate classic style resume - traditional, conservative format."""
        if format == 'pdf':
            return self._create_classic_resume_pdf(user_data)
        else:
            return self._create_classic_resume_docx(user_data)

    def _generate_minimal_resume(self, user_data: Dict, format: str) -> str:
        """Generate minimal style resume - clean, simple design."""
        if format == 'pdf':
            return self._create_minimal_resume_pdf(user_data)
        else:
            return self._create_minimal_resume_docx(user_data)

    def _create_classic_resume_pdf(self, user_data: Dict) -> str:
        """Create classic/traditional resume in PDF format."""
        import tempfile

        temp_dir = tempfile.gettempdir()
        filename = f"resume_classic_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        filepath = os.path.join(temp_dir, filename)

        doc = SimpleDocTemplate(filepath, pagesize=letter)
        story = []
        styles = getSampleStyleSheet()

        # Classic styles - traditional, serif fonts, black and white
        title_style = ParagraphStyle(
            'ClassicTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=colors.black,
            spaceAfter=3,
            alignment=1,  # Center
            fontName='Times-Bold'
        )

        heading_style = ParagraphStyle(
            'ClassicHeading',
            parent=styles['Heading2'],
            fontSize=12,
            textColor=colors.black,
            spaceAfter=6,
            spaceBefore=12,
            fontName='Times-Bold',
            underline=True
        )

        body_style = ParagraphStyle(
            'ClassicBody',
            parent=styles['Normal'],
            fontSize=10,
            fontName='Times-Roman'
        )

        # Personal Info
        personal = user_data.get('personalInfo', {})
        name = personal.get('fullName', 'Your Name')
        story.append(Paragraph(name.upper(), title_style))

        # Contact info in classic format
        contact_lines = []
        if personal.get('email'):
            contact_lines.append(personal['email'])
        if personal.get('phone'):
            contact_lines.append(personal['phone'])
        if personal.get('location'):
            location = personal['location']
            if isinstance(location, dict):
                loc_parts = []
                if location.get('city'): loc_parts.append(location['city'])
                if location.get('state'): loc_parts.append(location['state'])
                contact_lines.append(', '.join(loc_parts))
            else:
                contact_lines.append(str(location))

        for line in contact_lines:
            p = Paragraph(line, body_style)
            p.alignment = 1  # Center
            story.append(p)

        story.append(Spacer(1, 0.15*inch))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.black))
        story.append(Spacer(1, 0.1*inch))

        # Summary/Objective
        summary = user_data.get('summary', '')
        if summary:
            story.append(Paragraph('OBJECTIVE', heading_style))
            story.append(Paragraph(summary, body_style))
            story.append(Spacer(1, 0.15*inch))

        # Work Experience
        experiences = user_data.get('experiences', user_data.get('workExperience', []))
        if experiences:
            story.append(Paragraph('PROFESSIONAL EXPERIENCE', heading_style))
            for exp in experiences:
                title = exp.get('title', 'Position')
                company = exp.get('company', 'Company')
                start = exp.get('startDate', '')
                end = exp.get('endDate', 'Present')

                # Job title and dates on same line
                job_header = f"<b>{title}</b>, {company}"
                story.append(Paragraph(job_header, body_style))
                story.append(Paragraph(f"<i>{start} - {end}</i>", body_style))

                if exp.get('description'):
                    story.append(Paragraph(exp['description'], body_style))

                if exp.get('achievements'):
                    for achievement in exp['achievements']:
                        story.append(Paragraph(f"• {achievement}", body_style))

                story.append(Spacer(1, 0.1*inch))

        # Education
        education = user_data.get('education', [])
        if education:
            story.append(Paragraph('EDUCATION', heading_style))
            for edu in education:
                degree = edu.get('degree', 'Degree')
                institution = edu.get('institution', 'Institution')
                grad_date = edu.get('endDate', edu.get('graduationDate', ''))

                story.append(Paragraph(f"<b>{degree}</b>", body_style))
                story.append(Paragraph(f"{institution}, {grad_date}", body_style))

                if edu.get('gpa'):
                    story.append(Paragraph(f"GPA: {edu['gpa']}", body_style))

                story.append(Spacer(1, 0.1*inch))

        # Skills
        skills = user_data.get('skills', [])
        if skills:
            story.append(Paragraph('SKILLS', heading_style))
            # Format skills nicely
            if isinstance(skills[0], dict):
                skills_text = ', '.join([s.get('name', s.get('skill', '')) for s in skills[:20]])
            else:
                skills_text = ', '.join(skills[:20])
            story.append(Paragraph(skills_text, body_style))

        # Certifications
        certifications = user_data.get('certifications', [])
        if certifications:
            story.append(Spacer(1, 0.1*inch))
            story.append(Paragraph('CERTIFICATIONS', heading_style))
            for cert in certifications:
                if isinstance(cert, dict):
                    cert_name = cert.get('name', cert.get('certification', ''))
                    issuer = cert.get('issuer', '')
                    story.append(Paragraph(f"• {cert_name}" + (f" - {issuer}" if issuer else ""), body_style))
                else:
                    story.append(Paragraph(f"• {cert}", body_style))

        doc.build(story)
        return filepath

    def _create_classic_resume_docx(self, user_data: Dict) -> str:
        """Create classic resume in DOCX format."""
        import tempfile

        temp_dir = tempfile.gettempdir()
        filename = f"resume_classic_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(temp_dir, filename)

        doc = Document()

        # Personal Info - centered, traditional
        personal = user_data.get('personalInfo', {})
        name = personal.get('fullName', 'Your Name')

        name_para = doc.add_paragraph(name.upper())
        name_para.runs[0].bold = True
        name_para.runs[0].font.size = Pt(16)
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Contact info
        if personal.get('email'):
            p = doc.add_paragraph(personal['email'])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if personal.get('phone'):
            p = doc.add_paragraph(personal['phone'])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if personal.get('location'):
            location = personal['location']
            if isinstance(location, dict):
                loc_parts = []
                if location.get('city'): loc_parts.append(location['city'])
                if location.get('state'): loc_parts.append(location['state'])
                p = doc.add_paragraph(', '.join(loc_parts))
            else:
                p = doc.add_paragraph(str(location))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph('_' * 80)  # Horizontal line

        # Summary/Objective
        summary = user_data.get('summary', '')
        if summary:
            heading = doc.add_heading('OBJECTIVE', level=2)
            heading.runs[0].font.size = Pt(12)
            heading.runs[0].underline = True
            doc.add_paragraph(summary)

        # Work Experience
        experiences = user_data.get('experiences', user_data.get('workExperience', []))
        if experiences:
            heading = doc.add_heading('PROFESSIONAL EXPERIENCE', level=2)
            heading.runs[0].font.size = Pt(12)
            heading.runs[0].underline = True

            for exp in experiences:
                title = exp.get('title', 'Position')
                company = exp.get('company', 'Company')
                dates = f"{exp.get('startDate', '')} - {exp.get('endDate', 'Present')}"

                p = doc.add_paragraph()
                p.add_run(f"{title}, {company}").bold = True
                doc.add_paragraph(dates).italic = True

                if exp.get('description'):
                    doc.add_paragraph(exp['description'])

                if exp.get('achievements'):
                    for achievement in exp['achievements']:
                        doc.add_paragraph(f"• {achievement}", style='List Bullet')

        # Education
        education = user_data.get('education', [])
        if education:
            heading = doc.add_heading('EDUCATION', level=2)
            heading.runs[0].font.size = Pt(12)
            heading.runs[0].underline = True

            for edu in education:
                degree = edu.get('degree', 'Degree')
                institution = edu.get('institution', 'Institution')
                grad_date = edu.get('endDate', edu.get('graduationDate', ''))

                p = doc.add_paragraph()
                p.add_run(degree).bold = True
                doc.add_paragraph(f"{institution}, {grad_date}")

                if edu.get('gpa'):
                    doc.add_paragraph(f"GPA: {edu['gpa']}")

        # Skills
        skills = user_data.get('skills', [])
        if skills:
            heading = doc.add_heading('SKILLS', level=2)
            heading.runs[0].font.size = Pt(12)
            heading.runs[0].underline = True

            if isinstance(skills[0], dict):
                skills_text = ', '.join([s.get('name', s.get('skill', '')) for s in skills[:20]])
            else:
                skills_text = ', '.join(skills[:20])
            doc.add_paragraph(skills_text)

        # Certifications
        certifications = user_data.get('certifications', [])
        if certifications:
            heading = doc.add_heading('CERTIFICATIONS', level=2)
            heading.runs[0].font.size = Pt(12)
            heading.runs[0].underline = True

            for cert in certifications:
                if isinstance(cert, dict):
                    cert_name = cert.get('name', cert.get('certification', ''))
                    issuer = cert.get('issuer', '')
                    doc.add_paragraph(f"• {cert_name}" + (f" - {issuer}" if issuer else ""), style='List Bullet')
                else:
                    doc.add_paragraph(f"• {cert}", style='List Bullet')

        doc.save(filepath)
        return filepath

    def _create_minimal_resume_pdf(self, user_data: Dict) -> str:
        """Create minimal/clean resume in PDF format."""
        import tempfile

        temp_dir = tempfile.gettempdir()
        filename = f"resume_minimal_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        filepath = os.path.join(temp_dir, filename)

        doc = SimpleDocTemplate(filepath, pagesize=letter,
                              leftMargin=1*inch, rightMargin=1*inch,
                              topMargin=0.75*inch, bottomMargin=0.75*inch)
        story = []
        styles = getSampleStyleSheet()

        # Minimal styles - clean, lots of whitespace, subtle colors
        title_style = ParagraphStyle(
            'MinimalTitle',
            parent=styles['Heading1'],
            fontSize=28,
            textColor=colors.HexColor('#1a1a1a'),
            spaceAfter=4,
            fontName='Helvetica-Bold',
            leading=32
        )

        subtitle_style = ParagraphStyle(
            'MinimalSubtitle',
            parent=styles['Normal'],
            fontSize=10,
            textColor=colors.HexColor('#666666'),
            spaceAfter=20,
            fontName='Helvetica'
        )

        heading_style = ParagraphStyle(
            'MinimalHeading',
            parent=styles['Heading2'],
            fontSize=11,
            textColor=colors.HexColor('#333333'),
            spaceAfter=10,
            spaceBefore=20,
            fontName='Helvetica-Bold',
            letterSpacing=1
        )

        # Personal Info - minimal, left-aligned
        personal = user_data.get('personalInfo', {})
        name = personal.get('fullName', 'Your Name')
        story.append(Paragraph(name, title_style))

        # Contact - one line, minimal
        contact_parts = []
        if personal.get('email'):
            contact_parts.append(personal['email'])
        if personal.get('phone'):
            contact_parts.append(personal['phone'])
        if personal.get('location'):
            location = personal['location']
            if isinstance(location, dict):
                if location.get('city'):
                    contact_parts.append(location['city'])
            else:
                contact_parts.append(str(location))

        if contact_parts:
            story.append(Paragraph('  •  '.join(contact_parts), subtitle_style))

        # Summary - no label, just content
        summary = user_data.get('summary', '')
        if summary:
            story.append(Paragraph(summary, styles['Normal']))
            story.append(Spacer(1, 0.3*inch))

        # Experience - minimal formatting
        experiences = user_data.get('experiences', user_data.get('workExperience', []))
        if experiences:
            story.append(Paragraph('EXPERIENCE', heading_style))

            for exp in experiences:
                title = exp.get('title', 'Position')
                company = exp.get('company', 'Company')
                dates = f"{exp.get('startDate', '')} – {exp.get('endDate', 'Present')}"

                # Title and company
                job_style = ParagraphStyle('JobTitle', parent=styles['Normal'],
                                         fontSize=11, fontName='Helvetica-Bold')
                story.append(Paragraph(title, job_style))

                # Company and dates on same line
                company_style = ParagraphStyle('Company', parent=styles['Normal'],
                                             fontSize=10, textColor=colors.HexColor('#666666'))
                story.append(Paragraph(f"{company}  •  {dates}", company_style))

                if exp.get('description'):
                    story.append(Spacer(1, 0.05*inch))
                    story.append(Paragraph(exp['description'], styles['Normal']))

                if exp.get('achievements'):
                    for achievement in exp['achievements']:
                        story.append(Paragraph(f"• {achievement}", styles['Normal']))

                story.append(Spacer(1, 0.15*inch))

        # Education
        education = user_data.get('education', [])
        if education:
            story.append(Paragraph('EDUCATION', heading_style))

            for edu in education:
                degree = edu.get('degree', 'Degree')
                institution = edu.get('institution', 'Institution')
                grad_date = edu.get('endDate', edu.get('graduationDate', ''))

                degree_style = ParagraphStyle('Degree', parent=styles['Normal'],
                                            fontSize=11, fontName='Helvetica-Bold')
                story.append(Paragraph(degree, degree_style))

                school_style = ParagraphStyle('School', parent=styles['Normal'],
                                            fontSize=10, textColor=colors.HexColor('#666666'))
                story.append(Paragraph(f"{institution}  •  {grad_date}", school_style))

                story.append(Spacer(1, 0.1*inch))

        # Skills - clean list
        skills = user_data.get('skills', [])
        if skills:
            story.append(Paragraph('SKILLS', heading_style))

            if isinstance(skills[0], dict):
                skills_list = [s.get('name', s.get('skill', '')) for s in skills[:15]]
            else:
                skills_list = skills[:15]

            skills_text = '  •  '.join(skills_list)
            story.append(Paragraph(skills_text, styles['Normal']))

        doc.build(story)
        return filepath

    def _create_minimal_resume_docx(self, user_data: Dict) -> str:
        """Create minimal resume in DOCX format."""
        import tempfile

        temp_dir = tempfile.gettempdir()
        filename = f"resume_minimal_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(temp_dir, filename)

        doc = Document()

        # Set minimal margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Personal Info - clean, left-aligned
        personal = user_data.get('personalInfo', {})
        name = personal.get('fullName', 'Your Name')

        name_para = doc.add_paragraph(name)
        name_para.runs[0].bold = True
        name_para.runs[0].font.size = Pt(24)

        # Contact - one line
        contact_parts = []
        if personal.get('email'):
            contact_parts.append(personal['email'])
        if personal.get('phone'):
            contact_parts.append(personal['phone'])
        if personal.get('location'):
            location = personal['location']
            if isinstance(location, dict) and location.get('city'):
                contact_parts.append(location['city'])
            elif not isinstance(location, dict):
                contact_parts.append(str(location))

        if contact_parts:
            contact_para = doc.add_paragraph('  •  '.join(contact_parts))
            contact_para.runs[0].font.size = Pt(9)
            contact_para.runs[0].font.color.rgb = RGBColor(102, 102, 102)

        doc.add_paragraph()  # Spacer

        # Summary - no heading
        summary = user_data.get('summary', '')
        if summary:
            doc.add_paragraph(summary)
            doc.add_paragraph()  # Spacer

        # Experience
        experiences = user_data.get('experiences', user_data.get('workExperience', []))
        if experiences:
            heading = doc.add_paragraph('EXPERIENCE')
            heading.runs[0].bold = True
            heading.runs[0].font.size = Pt(11)

            for exp in experiences:
                title = exp.get('title', 'Position')
                company = exp.get('company', 'Company')
                dates = f"{exp.get('startDate', '')} – {exp.get('endDate', 'Present')}"

                # Title
                p = doc.add_paragraph(title)
                p.runs[0].bold = True
                p.runs[0].font.size = Pt(11)

                # Company and dates
                p = doc.add_paragraph(f"{company}  •  {dates}")
                p.runs[0].font.size = Pt(10)
                p.runs[0].font.color.rgb = RGBColor(102, 102, 102)

                if exp.get('description'):
                    doc.add_paragraph(exp['description'])

                if exp.get('achievements'):
                    for achievement in exp['achievements']:
                        doc.add_paragraph(f"• {achievement}")

        # Education
        education = user_data.get('education', [])
        if education:
            doc.add_paragraph()  # Spacer
            heading = doc.add_paragraph('EDUCATION')
            heading.runs[0].bold = True
            heading.runs[0].font.size = Pt(11)

            for edu in education:
                degree = edu.get('degree', 'Degree')
                institution = edu.get('institution', 'Institution')
                grad_date = edu.get('endDate', edu.get('graduationDate', ''))

                p = doc.add_paragraph(degree)
                p.runs[0].bold = True

                p = doc.add_paragraph(f"{institution}  •  {grad_date}")
                p.runs[0].font.size = Pt(10)
                p.runs[0].font.color.rgb = RGBColor(102, 102, 102)

        # Skills
        skills = user_data.get('skills', [])
        if skills:
            doc.add_paragraph()  # Spacer
            heading = doc.add_paragraph('SKILLS')
            heading.runs[0].bold = True
            heading.runs[0].font.size = Pt(11)

            if isinstance(skills[0], dict):
                skills_text = '  •  '.join([s.get('name', s.get('skill', '')) for s in skills[:15]])
            else:
                skills_text = '  •  '.join(skills[:15])
            doc.add_paragraph(skills_text)

        doc.save(filepath)
        return filepath
